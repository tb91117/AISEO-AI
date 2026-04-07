from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# Helper functions
def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x0F, 0x3F, 0x7F)

def body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)

def stage_note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

def timing(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x3F, 0x7F)

def divider():
    doc.add_paragraph("-" * 70)

# Title block
title = doc.add_heading("Loom Video Script", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("AEGIS Assignment Walkthrough  |  AI Founding Engineer - AISEO")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub2.add_run("Target runtime: 7-8 minutes")
r2.font.size = Pt(11)
r2.italic = True
r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
divider()
doc.add_paragraph()

# SECTION 1
heading1("SECTION 1 - Intro")
timing("0:00 - 0:30")
stage_note("Direction: Face on camera, no screen share yet")
doc.add_paragraph()
body(
    "\"Hi, I'm [YOUR NAME]. This is my walkthrough for the AEGIS take-home assignment "
    "for the AI Founding Engineer role at AISEO.\n\n"
    "I'll cover three things: how I approached the assignment, the technical decisions "
    "I made and why, and a quick overview of the experience I'm bringing to this role. "
    "Let me share my screen and walk through the code.\""
)

doc.add_paragraph()
divider()
doc.add_paragraph()

# SECTION 2
heading1("SECTION 2 - Project Overview")
timing("0:30 - 1:30")
stage_note("Direction: Switch to screen share - show the project folder structure")
doc.add_paragraph()
body(
    "\"The assignment asked me to build a Python FastAPI service called AEGIS with two features.\n\n"
    "Feature one is an AEO Content Scorer - it accepts a URL or raw HTML, runs three NLP checks, "
    "and returns a 0-to-100 readiness score. Feature two is a Query Fan-Out Engine - it calls an "
    "LLM to decompose a search query into 12 sub-queries across 6 intent types, then optionally "
    "runs semantic gap analysis against the user's existing content.\n\n"
    "Here's the full project structure.\""
)
doc.add_paragraph()
stage_note("Direction: Show file tree in terminal or editor")
doc.add_paragraph()
body(
    "\"Before writing a single line, I read the spec end to end and mapped each requirement to a "
    "file. The structure the assignment suggested maps cleanly to separation of concerns - checks "
    "are independently testable, the LLM logic is isolated from the endpoint, and the schema layer "
    "is separate from the business logic. That made the implementation straightforward.\""
)

doc.add_paragraph()
divider()
doc.add_paragraph()

# SECTION 3
heading1("SECTION 3 - Technical Walkthrough")
timing("1:30 - 5:30")

heading3("3a - Schemas   (1:30 - 2:00)")
stage_note("Direction: Open app/models/schemas.py")
doc.add_paragraph()
body(
    "\"I started with the Pydantic schemas because everything else depends on them. I modelled "
    "each check's output as a typed detail class - DirectAnswerDetails, HtagHierarchyDetails, "
    "ReadabilityDetails - and unified them under CheckResult using a Union type. This gives you "
    "full type safety and automatic OpenAPI documentation with zero extra work.\""
)

doc.add_paragraph()

heading3("3b - AEO Checks   (2:00 - 3:15)")
stage_note("Direction: Open app/services/aeo_checks/direct_answer.py")
doc.add_paragraph()
body(
    "\"The most interesting NLP decision was in Check A - Direct Answer Detection. The spec says "
    "to use spaCy's dependency parser to verify a declarative sentence. The subtlety here is that "
    "spaCy follows Universal Dependencies, where copular verbs like 'is' and 'are' are tagged as "
    "AUX, not VERB. A naive check for pos_ == VERB would incorrectly fail the sentence "
    "'Python is a high-level language' - which is a perfectly valid declarative answer. "
    "I caught this in testing and fixed it to accept both VERB and AUX at the ROOT dependency.\""
)
doc.add_paragraph()
stage_note("Direction: Open app/services/aeo_checks/readability.py")
doc.add_paragraph()
body(
    "\"For Check C - Readability, I used textstat's Flesch-Kincaid Grade Level. FK grade is a "
    "float, but the spec defines score bands in integer grades. I round to the nearest integer "
    "before scoring - so grade 9.4 stays in the 20-point band, but 9.6 maps to grade 10 and "
    "drops to 14 points. I also extract the top 3 most complex sentences by syllable density - "
    "syllable count divided by word count per sentence - which gives editors a concrete revision "
    "target.\""
)

doc.add_paragraph()

heading3("3c - Fan-Out Prompt   (3:15 - 4:15)")
stage_note("Direction: Open app/services/fanout_engine.py, scroll to _PROMPT_TEMPLATE")
doc.add_paragraph()
body(
    "\"The LLM prompt is what I spent the most iteration time on - it's documented in "
    "PROMPT_LOG.md. My first draft was a single sentence: generate 10-15 sub-queries, return "
    "as JSON. The model returned markdown-wrapped JSON with uneven type distribution - "
    "4 comparatives, 1 definitional.\n\n"
    "The two changes that fixed everything were: fixing the count to exactly 12, 2 per type "
    "instead of an open range, and adding a worked example using a completely different query "
    "topic. The different topic is important - if your example is about SEO tools and your test "
    "query is also about SEO tools, the model paraphrases the example instead of reasoning about "
    "the new query.\n\n"
    "I also added type definitions explaining the intent of each type, not just the name. That "
    "was the difference between the model generating a real use-case query versus outputting "
    "'a query about how to use the tool in a startup'.\""
)

doc.add_paragraph()

heading3("3d - Gap Analyzer   (4:15 - 4:45)")
stage_note("Direction: Open app/services/gap_analyzer.py")
doc.add_paragraph()
body(
    "\"For semantic gap analysis I chose all-MiniLM-L6-v2 over all-mpnet-base-v2. MiniLM is "
    "5 times faster on CPU and uses 80 MB instead of 420 MB - for a per-request operation "
    "encoding potentially hundreds of article sentences, that latency difference is user-facing. "
    "I encode both content chunks and sub-queries with normalize_embeddings=True so the dot "
    "product equals cosine similarity - avoiding the raw dot product on non-normalised vectors "
    "mistake flagged in the evaluation criteria.\""
)

doc.add_paragraph()

heading3("3e - Tests   (4:45 - 5:30)")
stage_note("Direction: Open tests folder, run  pytest tests/ -v  in terminal live")
doc.add_paragraph()
body(
    "\"The test suite covers 52 cases and runs entirely without an API key - all LLM calls are "
    "mocked. I tested both happy paths and error paths: invalid JSON from the LLM, missing keys, "
    "invalid type values, too few sub-queries, and all three retry attempts failing. The retry "
    "test uses patch('time.sleep') so it runs instantly.\""
)

doc.add_paragraph()
divider()
doc.add_paragraph()

# SECTION 4
heading1("SECTION 4 - Relevant Experience")
timing("5:30 - 7:00")
stage_note("Direction: Back to face camera, or stay on screen with GitHub / portfolio")
doc.add_paragraph()
body(
    "\"Let me briefly connect this to my background.\n\n"
    "[PLACEHOLDER - fill in from your resume, for example:]\n\n"
    "I've spent [X years] building production AI systems. Most recently at [COMPANY], I "
    "[built / led / architected] [SPECIFIC SYSTEM] that [QUANTIFIED OUTCOME - e.g. processed "
    "10M requests/day, reduced latency by 40%].\n\n"
    "On the LLM side, I have hands-on experience with [prompt engineering / fine-tuning / RAG "
    "systems] - specifically [SPECIFIC PROJECT OR TOOL]. The prompt engineering work in this "
    "assignment reflects how I approach it in production: iterate on a real failure, document "
    "why each change was made, and validate systematically.\n\n"
    "On the SEO and GEO side, [RELEVANT EXPERIENCE if any].\n\n"
    "What excites me about AISEO specifically is the GEO angle. The shift from ten blue links "
    "to AI-generated answers is the most significant change in search since mobile - and building "
    "the tooling that helps content teams navigate that change is exactly the kind of "
    "high-leverage technical work I want to be doing.\""
)

doc.add_paragraph()
divider()
doc.add_paragraph()

# SECTION 5
heading1("SECTION 5 - Close")
timing("7:00 - 7:30")
stage_note("Direction: Face on camera")
doc.add_paragraph()
body(
    "\"To summarise: Feature 1 is fully complete with 3 NLP checks, typed schemas, and unit "
    "tests. Feature 2 is complete with a well-engineered prompt, retry logic, and semantic gap "
    "analysis. The prompt log and README cover all the design decisions.\n\n"
    "I'm happy to go deeper on any of these decisions in the follow-up interview. Thanks for "
    "the opportunity - looking forward to the conversation.\""
)

doc.add_paragraph()
divider()
doc.add_paragraph()

# Recording Tips
heading1("Recording Tips")
tips = [
    "Screen layout: VS Code on the left, terminal on the right. Dark theme reads better on Loom.",
    "Don't read the script verbatim - use it as bullet points and talk naturally.",
    "Run  pytest tests/ -v  live so the reviewer sees all 52 tests pass in real time.",
    "Show the Swagger UI at /docs and make one live API call to /api/aeo/analyze - more compelling than static code.",
    "Keep each section tight; finishing at 7 minutes is better than padding to 10.",
]
for tip in tips:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(tip)
    run.font.size = Pt(11)

# Save
out = r"E:\Work\Example-Projects\_Test_Employment\AI-founding-engineer-assignment\Loom_Script_AEGIS.docx"
doc.save(out)
print(f"Saved: {out}")
